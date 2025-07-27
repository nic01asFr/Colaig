"""
Module pour l'analyse des documents.
Fournit des fonctionnalités pour extraire du texte et analyser des documents.
"""

import os
import tempfile
import asyncio
import io
from typing import Optional, Any, Dict, BinaryIO, List

from app.matrix_bot.config import logger

from app.core_llm import AlbertApiClient

# Imports pour l'extraction de texte des PDFs
try:
    from pdfminer.high_level import extract_text as pdfminer_extract_text
    from pdfminer.layout import LAParams, LTTextContainer, LTChar
    from pdfminer.pdfpage import PDFPage
    from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
    from pdfminer.converter import PDFPageAggregator
    PDFMINER_AVAILABLE = True
except ImportError:
    logger.warning("[DOC_ANALYZER] pdfminer.six n'est pas disponible, l'extraction de texte PDF sera limitée")
    PDFMINER_AVAILABLE = False

# Tentative d'import des bibliothèques pour d'autres formats
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    logger.warning("[DOC_ANALYZER] python-docx n'est pas disponible, l'extraction DOCX sera limitée")
    DOCX_AVAILABLE = False


class DocumentAnalyzer:
    """
    Classe pour analyser des documents.
    Extrait le texte et génère des résumés de documents.
    """
    
    def __init__(self, config: Any):
        """
        Initialise l'analyseur de documents.
        
        Args:
            config: Configuration du bot contenant les clés API
        """
        self.config = config
    
    async def extract_text(self, file: BinaryIO, filename: str) -> str:
        """
        Extrait le texte d'un document.
        
        Args:
            file: Fichier binaire à analyser
            filename: Nom du fichier
            
        Returns:
            Le texte extrait du document
        """
        # Créer un fichier temporaire
        temp_fd, temp_path = tempfile.mkstemp(suffix=os.path.splitext(filename)[1])
        
        try:
            # Écrire le contenu dans le fichier temporaire
            with os.fdopen(temp_fd, 'wb') as temp_file:
                file.seek(0)  # S'assurer d'être au début du fichier
                file_content = file.read()
                temp_file.write(file_content)
            
            # Extraire le texte selon le type de fichier
            file_ext = os.path.splitext(filename)[1].lower()
            
            # Extraction basique pour les fichiers texte
            if file_ext in ['.txt', '.md', '.json', '.xml', '.csv']:
                with open(temp_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            
            # Extraction pour les fichiers PDF
            elif file_ext == '.pdf' and PDFMINER_AVAILABLE:
                logger.info(f"[DOC_ANALYZER] Extraction du texte du PDF: {filename}")
                return await self._extract_pdf_text(io.BytesIO(file_content), filename)
            
            # Extraction pour les fichiers DOCX
            elif file_ext in ['.docx', '.doc'] and DOCX_AVAILABLE:
                logger.info(f"[DOC_ANALYZER] Extraction du texte du DOCX: {filename}")
                return self._extract_docx_text(temp_path, filename)
            
            # Pour les autres types de fichiers, on renvoie un message d'information
            return f"Le document \"{filename}\" est un fichier {file_ext.upper()[1:]}. Il n'est pas possible d'en extraire le texte de manière avancée pour le moment. Pour plus d'informations, il serait nécessaire d'accéder directement au contenu du fichier."
            
        except Exception as e:
            logger.error(f"[DOC_ANALYZER] Erreur lors de l'extraction de texte: {str(e)}")
            return f"Impossible d'extraire le texte du document '{filename}': {str(e)}"
        
        finally:
            # Nettoyer
            try:
                os.remove(temp_path)
            except:
                pass

    async def _extract_pdf_text(self, file_content: BinaryIO, filename: str) -> str:
        """
        Extrait le texte d'un fichier PDF en utilisant pdfminer.six
        
        Args:
            file_content: Contenu binaire du fichier PDF
            filename: Nom du fichier pour les logs
            
        Returns:
            Le texte extrait du PDF
        """
        try:
            # Méthode 1: Utiliser extract_text (simple mais moins de contrôle)
            try:
                file_content.seek(0)
                simple_text = pdfminer_extract_text(file_content)
                
                # Si extract_text a fonctionné et donné un texte non vide, l'utiliser
                if simple_text and len(simple_text.strip()) > 50:
                    logger.info(f"[DOC_ANALYZER] Texte extrait avec extract_text: {len(simple_text)} caractères")
                    return simple_text
            except Exception as simple_err:
                logger.warning(f"[DOC_ANALYZER] Erreur avec extract_text: {str(simple_err)}, tentative avec méthode avancée")
            
            # Méthode 2: Utiliser PDFMiner avec plus de contrôle
            # Paramètres d'extraction améliorés
            laparams = LAParams(
                line_margin=0.3,      # Réduit pour mieux détecter les lignes
                word_margin=0.1,      # Marge entre les mots
                char_margin=1.0,      # Réduit pour mieux grouper les caractères
                boxes_flow=0.7,       # Augmenté pour mieux suivre le flux de texte
                detect_vertical=True,  # Détection du texte vertical
                all_texts=True        # Capture tout le texte, même décoratif
            )
            
            file_content.seek(0)
            
            structured_text = []
            current_page = 1
            
            # Extraire le texte page par page
            for page in PDFPage.get_pages(file_content):
                # Initialiser les outils PDFMiner pour chaque page
                resource_manager = PDFResourceManager()
                device = PDFPageAggregator(resource_manager, laparams=laparams)
                interpreter = PDFPageInterpreter(resource_manager, device)
                
                # Traiter la page
                interpreter.process_page(page)
                layout = device.get_result()
                
                page_text = []
                for element in layout:
                    if isinstance(element, LTTextContainer):
                        text = element.get_text().strip()
                        if text:
                            page_text.append(text)
                
                # Ajouter le texte de la page avec un délimitateur
                if page_text:
                    structured_text.append(f"[Page {current_page}]\n" + "\n".join(page_text))
                
                # Nettoyer les ressources
                device.close()
                current_page += 1
            
            # Joindre tout le texte
            full_text = "\n\n".join(structured_text)
            logger.info(f"[DOC_ANALYZER] Texte structuré extrait: {len(full_text)} caractères, {current_page-1} pages")
            
            # Si la méthode avancée n'a pas donné de résultats, essayer extract_text à nouveau
            if not full_text or len(full_text.strip()) < 10:
                file_content.seek(0)
                fallback_text = pdfminer_extract_text(file_content)
                if fallback_text and len(fallback_text.strip()) > 0:
                    logger.info(f"[DOC_ANALYZER] Texte extrait par fallback: {len(fallback_text)} caractères")
                    return fallback_text
                
                # Si toujours pas de texte, c'est probablement un PDF numérisé ou protégé
                if not fallback_text or len(fallback_text.strip()) < 10:
                    return f"Le document \"{filename}\" est un fichier PDF qui semble être numérisé ou protégé. L'extraction de texte n'a pas donné de résultats."
            
            return full_text
            
        except Exception as e:
            logger.error(f"[DOC_ANALYZER] Erreur lors de l'extraction PDF: {str(e)}")
            return f"Le document \"{filename}\" est un fichier PDF. Une erreur s'est produite lors de l'extraction: {str(e)}"
    
    def _extract_docx_text(self, file_path: str, filename: str) -> str:
        """
        Extrait le texte d'un fichier DOCX en utilisant python-docx
        
        Args:
            file_path: Chemin vers le fichier temporaire
            filename: Nom du fichier pour les logs
            
        Returns:
            Le texte extrait du DOCX
        """
        try:
            if DOCX_AVAILABLE:
                # Ouvrir le document
                doc = docx.Document(file_path)
                
                # Extraire le texte de chaque paragraphe
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                
                # Extraire le texte des tableaux
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            full_text.append(cell.text)
                
                # Joindre tout le texte
                result = "\n".join(full_text)
                logger.info(f"[DOC_ANALYZER] Texte extrait du DOCX: {len(result)} caractères")
                return result
            else:
                return f"Le document \"{filename}\" est un fichier DOCX. L'extraction de texte n'est pas disponible car la bibliothèque python-docx n'est pas installée."
        except Exception as e:
            logger.error(f"[DOC_ANALYZER] Erreur lors de l'extraction DOCX: {str(e)}")
            return f"Le document \"{filename}\" est un fichier DOCX. Une erreur s'est produite lors de l'extraction: {str(e)}"
    
    async def analyze_document(self, file: BinaryIO, filename: str) -> str:
        """
        Analyse un document et génère un résumé.
        
        Args:
            file: Fichier binaire à analyser
            filename: Nom du fichier
            
        Returns:
            Un résumé du document
        """
        logger.info(f"[DOC_ANALYZER] Analyse du document '{filename}'")
        
        try:
            # Extraire le texte
            text = await self.extract_text(file, filename)
            
            # Si pas de texte ou texte trop court, renvoyer un message d'erreur
            if not text or len(text.strip()) < 50:
                return f"Le document '{filename}' semble être vide ou ne contient pas assez de texte pour être analysé."
            
            # Tronquer le texte si trop long
            max_text_length = 8000  # Limite raisonnable pour éviter les problèmes de contexte
            if len(text) > max_text_length:
                logger.info(f"[DOC_ANALYZER] Document '{filename}' tronqué de {len(text)} à {max_text_length} caractères")
                text = text[:max_text_length] + "...[texte tronqué]"
            
            # Générer le résumé avec le LLM
            api_key = self.config.albert_api_token
            url = self.config.albert_api_url
            model = self.config.albert_model
            
            # Créer le client
            client = AlbertApiClient(base_url=url, api_key=api_key)
            
            # Construire le prompt
            system_prompt = """Tu es Albert, un assistant d'analyse de documents. Tu dois analyser le document fourni et produire un résumé structuré.
Ton analyse doit inclure:
1. Un résumé du contenu principal du document
2. Les points clés ou informations importantes
3. Une structure ou organisation du document si pertinent
4. Le type et le style du document

Sois factuel et précis. N'invente pas d'informations qui ne sont pas présentes dans le document."""
            
            # Construire les messages
            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"""Voici le contenu du document '{filename}' à analyser:

{text}

Merci de fournir une analyse structurée de ce document."""}
            ]
            
            # Générer la réponse
            logger.info(f"[DOC_ANALYZER] Génération du résumé pour '{filename}'")
            response = await client.generate(
                model=model,
                messages=messages,
                temperature=0.3,
                max_tokens=1000
            )
            
            # Fermer le client
            await client.close()
            
            logger.info(f"[DOC_ANALYZER] Résumé généré avec succès pour '{filename}'")
            return response
            
        except Exception as e:
            logger.error(f"[DOC_ANALYZER] Erreur lors de l'analyse du document: {str(e)}")
            return f"Impossible d'analyser le document '{filename}': {str(e)}" 